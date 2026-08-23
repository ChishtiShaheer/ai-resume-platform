"""
Extracts raw text from an uploaded resume file (PDF or DOCX).

Strategy for PDFs:
  1. Try direct text extraction with pdfplumber (fast, works for
     text-based/"born-digital" PDFs).
  2. If the extracted text is too short (a strong signal the PDF is a
     scanned image with no text layer), fall back to OCR.

DOCX files are parsed directly via python-docx (paragraphs + tables).
"""
import logging
from dataclasses import dataclass

import pdfplumber
import docx

from app.services.ocr_service import ocr_pdf, is_ocr_available

logger = logging.getLogger(__name__)

MIN_TEXT_LENGTH_BEFORE_OCR = 40  # chars; below this we assume a scanned PDF


@dataclass
class ParseResult:
    text: str
    used_ocr: bool
    page_count: int = 0


def _extract_pdf_text(file_path: str) -> ParseResult:
    text_chunks = []
    page_count = 0
    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
            # Tables often hold skills/education in structured resumes
            for table in page.extract_tables() or []:
                for row in table:
                    text_chunks.append(" ".join([c for c in row if c]))
    text = "\n".join(text_chunks).strip()

    if len(text) < MIN_TEXT_LENGTH_BEFORE_OCR and is_ocr_available():
        logger.info("PDF %s looks scanned (only %d chars) — running OCR", file_path, len(text))
        ocr_text = ocr_pdf(file_path)
        if ocr_text and len(ocr_text) > len(text):
            return ParseResult(text=ocr_text, used_ocr=True, page_count=page_count)

    return ParseResult(text=text, used_ocr=False, page_count=page_count)


def _extract_docx_text(file_path: str) -> ParseResult:
    document = docx.Document(file_path)
    chunks = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            chunks.append(" ".join(cell.text for cell in row.cells if cell.text))
    return ParseResult(text="\n".join(chunks).strip(), used_ocr=False)


def parse_resume(file_path: str, file_type: str) -> ParseResult:
    file_type = file_type.lower()
    if file_type == "pdf":
        return _extract_pdf_text(file_path)
    elif file_type == "docx":
        return _extract_docx_text(file_path)
    raise ValueError(f"Unsupported file type: {file_type}")
